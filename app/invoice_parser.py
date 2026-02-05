"""
Invoice file parser for PDF, DOC, TXT, and XML files
Extracts invoice information and line items
"""

import re
from typing import Dict, List, Optional, Tuple
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def parse_pdf(file_path: str) -> Dict:
    """Parse PDF invoice file"""
    try:
        import PyPDF2
        
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        return extract_invoice_data(text, "PDF")
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        return {"error": str(e), "raw_text": ""}


def parse_docx(file_path: str) -> Dict:
    """Parse DOCX invoice file"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + "\t".join([cell.text for cell in row.cells])
        
        return extract_invoice_data(text, "DOCX")
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}")
        return {"error": str(e), "raw_text": ""}


def parse_txt(file_path: str) -> Dict:
    """Parse TXT invoice file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        return extract_invoice_data(text, "TXT")
    except Exception as e:
        logger.error(f"Error parsing TXT: {e}")
        return {"error": str(e), "raw_text": ""}


def parse_xml(file_path: str) -> Dict:
    """Parse XML invoice file (e-factura format or generic XML)"""
    try:
        from lxml import etree
        
        # First, check if file exists and is not empty
        path_obj = Path(file_path)
        if not path_obj.exists():
            return {"error": "XML file not found"}
        
        file_size = path_obj.stat().st_size
        if file_size == 0:
            return {"error": "XML file is empty"}
        
        # Read and strip BOM if present
        file_content = path_obj.read_text(encoding='utf-8-sig')
        if not file_content.strip():
            return {"error": "XML file contains no data"}
        
        # Parse XML from string (handles BOM better than parse from file)
        root = etree.fromstring(file_content.encode('utf-8'))
        
        # Try to extract common fields
        result = {
            "supplier": None,
            "invoice_number": None,
            "invoice_date": None,
            "total_amount": None,
            "items": [],
            "raw_text": etree.tostring(root, pretty_print=True).decode('utf-8')
        }
        
        # Define namespaces for e-factura (UBL format) to prevent "invalid predicate" errors
        namespaces = {
            'cac': 'urn:oasis:names:specification:ubl:schema:xsd:CommonAggregateComponents-2',
            'cbc': 'urn:oasis:names:specification:ubl:schema:xsd:CommonBasicComponents-2',
            'ubl': 'urn:oasis:names:specification:ubl:schema:xsd:Invoice-2'
        }
        
        def find_with_fallback(root_elem, xpath_with_ns, xpath_without_ns=None):
            """Try to find element with namespaces, fallback to without. Prevents 'invalid predicate' errors."""
            try:
                elem = root_elem.find(xpath_with_ns, namespaces)
                if elem is not None:
                    return elem
            except Exception:
                pass  # XPath failed, try fallback
            # Fallback: try without namespaces
            if xpath_without_ns:
                try:
                    return root_elem.find(xpath_without_ns)
                except Exception:
                    pass
            return None
        
        # Try UBL e-factura format first
        supplier_elem = find_with_fallback(
            root,
            './/cac:AccountingSupplierParty/cac:Party/cac:PartyName/cbc:Name',
            './/Party/PartyName/Name'
        )
        if supplier_elem is not None:
            result["supplier"] = supplier_elem.text
        
        invoice_number_elem = find_with_fallback(
            root,
            './/cbc:ID',
            './/ID'
        )
        if invoice_number_elem is not None:
            result["invoice_number"] = invoice_number_elem.text
        
        date_elem = find_with_fallback(
            root,
            './/cbc:IssueDate',
            './/IssueDate'
        )
        if date_elem is not None:
            result["invoice_date"] = date_elem.text
        
        total_elem = find_with_fallback(
            root,
            './/cac:LegalMonetaryTotal/cbc:TaxInclusiveAmount',
            './/TaxInclusiveAmount'
        )
        if total_elem is None:
            total_elem = find_with_fallback(
                root,
                './/cac:LegalMonetaryTotal/cbc:PayableAmount',
                './/PayableAmount'
            )
        if total_elem is not None and total_elem.text:
            try:
                result["total_amount"] = float(total_elem.text.replace(',', '.'))
            except:
                pass
        
        # Generic fallback XML parsing - adapt based on your specific XML format
        # This runs if UBL extraction didn't find everything
        if not result["supplier"] or not result["invoice_number"]:
            for elem in root.iter():
                text = (elem.text or "").strip()
                if not text:
                    continue
                    
                # Try to identify fields
                if not result["supplier"] and ('supplier' in elem.tag.lower() or 'furnizor' in elem.tag.lower() or 'party' in elem.tag.lower() and 'name' in elem.tag.lower()):
                    result["supplier"] = text
                elif not result["invoice_number"] and ('invoice' in elem.tag.lower() or 'factura' in elem.tag.lower()) and ('number' in elem.tag.lower() or 'id' in elem.tag.lower() or 'nr' in elem.tag.lower()):
                    result["invoice_number"] = text
                elif not result["invoice_date"] and ('date' in elem.tag.lower() or 'data' in elem.tag.lower() or 'issue' in elem.tag.lower()):
                    result["invoice_date"] = text
                elif not result["total_amount"] and 'total' in elem.tag.lower():
                    try:
                        result["total_amount"] = float(text.replace(',', '.'))
                    except:
                        pass
        
        # Try to extract items from UBL format first
        items = []
        try:
            invoice_lines = root.findall('.//cac:InvoiceLine', namespaces)
            if not invoice_lines:
                invoice_lines = root.findall('.//InvoiceLine')
            
            for line in invoice_lines:
                item = {
                    "description": None,
                    "quantity": 0.0,
                    "unit": None,
                    "unit_price": None,
                    "total_price": None
                }
                
                # Extract description
                desc_elem = find_with_fallback(line, './/cac:Item/cbc:Description', './/Item/Description')
                if desc_elem is None:
                    desc_elem = find_with_fallback(line, './/cac:Item/cbc:Name', './/Item/Name')
                if desc_elem is not None:
                    item["description"] = desc_elem.text
                
                # Extract quantity
                qty_elem = find_with_fallback(line, './/cbc:InvoicedQuantity', './/InvoicedQuantity')
                if qty_elem is not None:
                    try:
                        item["quantity"] = float(qty_elem.text.replace(',', '.'))
                        # Get unit from attribute
                        unit_code = qty_elem.get('unitCode') or qty_elem.get('unit')
                        if unit_code:
                            item["unit"] = unit_code
                    except:
                        pass
                
                # Extract unit price
                price_elem = find_with_fallback(line, './/cac:Price/cbc:PriceAmount', './/Price/PriceAmount')
                if price_elem is not None:
                    try:
                        item["unit_price"] = float(price_elem.text.replace(',', '.'))
                    except:
                        pass
                
                # Extract line total
                total_elem = find_with_fallback(line, './/cbc:LineExtensionAmount', './/LineExtensionAmount')
                if total_elem is not None:
                    try:
                        item["total_price"] = float(total_elem.text.replace(',', '.'))
                    except:
                        pass
                
                if item["description"]:
                    items.append(item)
        except Exception:
            pass  # Fall through to generic extraction
        
        # Generic fallback for non-UBL XML
        if not items:
            for item_elem in root.findall(".//*[contains(local-name(), 'item') or contains(local-name(), 'line')]"):
                item = {
                    "description": None,
                    "quantity": 0.0,
                    "unit": None,
                    "unit_price": None,
                    "total_price": None
                }
                
                for child in item_elem:
                    text = (child.text or "").strip()
                    if not text:
                        continue
                    if 'desc' in child.tag.lower() or 'name' in child.tag.lower():
                        item["description"] = text
                    elif 'quant' in child.tag.lower() or 'qty' in child.tag.lower():
                        try:
                            item["quantity"] = float(text.replace(',', '.'))
                            # Check for unit in attributes
                            unit_attr = child.get('unitCode') or child.get('unit')
                            if unit_attr:
                                item["unit"] = unit_attr
                        except:
                            pass
                    elif 'unit' in child.tag.lower() and 'price' not in child.tag.lower():
                        item["unit"] = text
                    elif 'price' in child.tag.lower() and 'total' not in child.tag.lower():
                        try:
                            item["unit_price"] = float(text.replace(',', '.'))
                        except:
                            pass
                    elif 'total' in child.tag.lower() or 'amount' in child.tag.lower():
                        try:
                            item["total_price"] = float(text.replace(',', '.'))
                        except:
                            pass
                
                if item["description"]:
                    items.append(item)
        
        result["items"] = items
        return result
        
    except Exception as e:
        logger.error(f"Error parsing XML: {e}")
        error_msg = str(e)
        # Provide more helpful error messages
        if "empty" in error_msg.lower():
            error_msg = "XML file is empty or contains no valid data"
        elif "line 1, column 1" in error_msg:
            error_msg = "XML file appears to be empty or has invalid encoding. Please ensure the file is a valid XML document."
        return {"error": error_msg, "raw_text": ""}


def extract_invoice_data(text: str, file_type: str) -> Dict:
    """
    Extract invoice information from text using regex patterns
    Returns dict with supplier, invoice_number, date, total, and items
    """
    result = {
        "supplier": None,
        "invoice_number": None,
        "invoice_date": None,
        "total_amount": None,
        "items": [],
        "raw_text": text
    }
    
    # Extract supplier (look for common patterns)
    supplier_patterns = [
        r'(?:supplier|furnizor|from)[:\s]+([^\n]+)',
        r'(?:company|societate)[:\s]+([^\n]+)',
    ]
    for pattern in supplier_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            result["supplier"] = match.group(1).strip()
            break
    
    # Extract invoice number
    invoice_patterns = [
        r'(?:invoice|factura)[^:]*(?:no|nr|number|numar)[.:\s#]+([A-Z0-9\-]+)',
        r'(?:no|nr)[.:\s]+([A-Z0-9\-]+)',
    ]
    for pattern in invoice_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            result["invoice_number"] = match.group(1).strip()
            break
    
    # Extract date (various formats)
    date_patterns = [
        r'(?:date|data)[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
        r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
        r'(\d{4}-\d{2}-\d{2})',
    ]
    for pattern in date_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            result["invoice_date"] = match.group(1).strip()
            break
    
    # Extract total amount
    total_patterns = [
        r'(?:total|total general)[:\s]+([0-9.,]+)',
        r'(?:total)[^0-9]+([0-9.,]+)',
    ]
    for pattern in total_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                total_str = match.group(1).replace(',', '.')
                result["total_amount"] = float(total_str)
            except:
                pass
            break
    
    # Extract line items (this is more complex and may need customization)
    result["items"] = extract_line_items(text)
    
    return result


def extract_line_items(text: str) -> List[Dict]:
    """
    Extract line items from invoice text
    Enhanced to handle ROMSTAL format with SKU codes and filtering
    """
    items = []
    
    # Look for table-like structures
    lines = text.split('\n')
    
    # Define patterns to skip (non-product lines)
    skip_patterns = [
        r'^\s*(BANCA|CONT|Capital|Sediul|Nr\.ord|C\.I\.F|Numar|TOTAL|Semnatura|Expedierea)',
        r'^\s*Date privind',
        r'^\s*(Cumparator|Furnizor):',
        r'^\s*Scadenta:',
        r'^\s*GESTIUNEA',
        r'^\s*Total.*:',
        r'^\s*Semnatura',
        r'^\s*Mijloc de transport',
        r'^\s*Numele delegatului',
        r'^\s*Buletin',
        r'^\s*Emis de'
    ]
    
    # Enhanced pattern for ROMSTAL format: 
    # Line number + SKU + Description + Unit + Quantity + Unit Price + Total
    # Example: "1 35FV1598 +Invertor monofazat... buc 1,000 3.179,84 3.179,84"
    romstal_pattern = re.compile(
        r'^\d+\s+([A-Z0-9]+)\s+(.+?)\s+(buc|kg|m|l|h|set|pcs|bucăți)\s+([\d.,]+)\s+([\d.,]+)\s+([\d.,]+)',
        re.IGNORECASE
    )
    
    # Standard pattern with optional unit
    standard_pattern = re.compile(
        r'(.+?)\s+([\d.,]+)\s+(?:buc|bucăți|kg|m|l|h|set|pcs)?\s*([\d.,]+)\s+([\d.,]+)',
        re.IGNORECASE
    )
    
    for line in lines:
        line = line.strip()
        if not line or len(line) < 10:
            continue
        
        # Skip non-product lines
        skip_line = False
        for pattern in skip_patterns:
            if re.search(pattern, line, re.IGNORECASE):
                skip_line = True
                break
        if skip_line:
            continue
        
        # Try ROMSTAL format first
        match = romstal_pattern.search(line)
        if match:
            try:
                sku = match.group(1).strip()
                description = match.group(2).strip()
                unit = match.group(3).strip()
                quantity_str = match.group(4).replace('.', '').replace(',', '.')
                unit_price_str = match.group(5).replace('.', '').replace(',', '.')
                total_str = match.group(6).replace('.', '').replace(',', '.')
                
                # Clean description - remove leading special characters
                description = re.sub(r'^[\+\-@\*#&\s]+', '', description)  # Remove leading special chars
                description = re.sub(r'[\+\-@\*#&\s]+$', '', description)  # Remove trailing special chars
                description = re.sub(r'\s+', ' ', description)  # Remove extra spaces
                description = re.sub(r'\bRON\b', '', description, flags=re.IGNORECASE)  # Remove RON
                description = description.strip()
                
                # Skip if description is too short or just unit
                if len(description) < 3 or description.lower() in ['buc', 'kg', 'm', 'l', 'h', 'set', 'pcs']:
                    continue
                
                # Add SKU to description if both exist
                if sku and description:
                    full_description = f"{sku} - {description}"
                elif sku:
                    full_description = sku
                else:
                    full_description = description
                
                quantity = float(quantity_str)
                unit_price = float(unit_price_str)
                total_price = float(total_str)
                
                items.append({
                    "description": full_description,
                    "quantity": quantity,
                    "unit": unit,
                    "unit_price": unit_price,
                    "total_price": total_price
                })
                continue
            except Exception as e:
                # Log error but continue to next line
                logger.debug(f"Failed to parse ROMSTAL line: {e}")
        
        # Try standard pattern
        match = standard_pattern.search(line)
        if match:
            try:
                description = match.group(1).strip()
                quantity_str = match.group(2).replace('.', '').replace(',', '.')
                unit_price_str = match.group(3).replace('.', '').replace(',', '.')
                total_str = match.group(4).replace('.', '').replace(',', '.')
                
                # Clean description - remove currency symbols and extra spaces
                description = re.sub(r'\bRON\b', '', description, flags=re.IGNORECASE)
                description = re.sub(r'\s+', ' ', description)
                description = description.strip()
                
                # Skip if description is just "buc" or other unit
                if description.lower() in ['buc', 'kg', 'm', 'l', 'h', 'set', 'pcs', 'bucăți']:
                    continue
                
                quantity = float(quantity_str)
                unit_price = float(unit_price_str)
                total_price = float(total_str)
                
                # Extract unit from description if present
                unit = None
                unit_match = re.search(r'\b(buc|bucăți|kg|m|l|h|set|pcs)\b', description, re.IGNORECASE)
                if unit_match:
                    unit = unit_match.group(1)
                    # Remove unit from description
                    description = description.replace(unit_match.group(0), '').strip()
                
                # Basic validation
                if len(description) > 2 and (abs(quantity * unit_price - total_price) < 1.0 or total_price > 0):
                    items.append({
                        "description": description,
                        "quantity": quantity,
                        "unit": unit,
                        "unit_price": unit_price,
                        "total_price": total_price
                    })
            except Exception as e:
                logger.debug(f"Failed to parse standard line: {e}")
                continue
    
    # If no items found, try simpler pattern (just description and numbers)
    if not items:
        simple_pattern = re.compile(r'(.{10,})\s+([\d.,]+)\s+([\d.,]+)', re.IGNORECASE)
        for line in lines:
            line = line.strip()
            match = simple_pattern.search(line)
            if match:
                try:
                    description = match.group(1).strip()
                    # Clean description
                    description = re.sub(r'\bRON\b', '', description, flags=re.IGNORECASE)
                    description = re.sub(r'\s+', ' ', description).strip()
                    
                    if len(description) < 3:
                        continue
                    
                    quantity_str = match.group(2).replace('.', '').replace(',', '.')
                    unit_price_str = match.group(3).replace('.', '').replace(',', '.')
                    
                    items.append({
                        "description": description,
                        "quantity": float(quantity_str),
                        "unit": None,
                        "unit_price": float(unit_price_str),
                        "total_price": None
                    })
                except:
                    continue
    
    return items


def parse_invoice_file(file_path: str) -> Dict:
    """
    Main function to parse invoice file based on extension
    """
    path = Path(file_path)
    extension = path.suffix.lower()
    
    if extension == '.pdf':
        return parse_pdf(file_path)
    elif extension in ['.doc', '.docx']:
        return parse_docx(file_path)
    elif extension == '.txt':
        return parse_txt(file_path)
    elif extension == '.xml':
        return parse_xml(file_path)
    else:
        return {"error": f"Unsupported file type: {extension}", "raw_text": ""}


def fuzzy_match_material(description: str, existing_materials: List[Dict]) -> Optional[Tuple[int, float]]:
    """
    Find the best matching material from existing materials
    Returns (material_id, confidence_score) or None
    """
    if not description or not existing_materials:
        return None
    
    description_lower = description.lower()
    best_match = None
    best_score = 0.0
    
    for material in existing_materials:
        material_name = (material.get('name') or '').lower()
        
        # Simple similarity scoring
        # Check if material name is in description
        if material_name in description_lower:
            score = len(material_name) / len(description_lower)
            if score > best_score:
                best_score = score
                best_match = material.get('id')
        
        # Check if description is in material name
        elif description_lower in material_name:
            score = len(description_lower) / len(material_name)
            if score > best_score:
                best_score = score
                best_match = material.get('id')
    
    # Only return if confidence is reasonable
    if best_match and best_score > 0.3:
        return (best_match, best_score)
    
    return None


# ============================================================================
# CSV Parsing and Enhanced Fuzzy Matching Functions
# ============================================================================

import csv
import io
from difflib import get_close_matches, SequenceMatcher

def normalize_string(s: str) -> str:
    """Normalize string for better matching"""
    if not s:
        return ""
    # Convert to lowercase, remove extra spaces, strip
    return " ".join(s.lower().strip().split())


def match_material_fuzzy(description: str, sku: Optional[str], materials: List[Dict], cutoff: float = 0.6) -> Optional[Dict]:
    """
    Match invoice line to material in database using difflib fuzzy matching
    
    Args:
        description: Item description from invoice
        sku: Optional SKU/code from invoice
        materials: List of materials from DB (dict with 'id', 'name', 'sku')
        cutoff: Minimum similarity score (0.0-1.0)
    
    Returns:
        Dict with matched material and confidence score, or None
    """
    if not description and not sku:
        return None
    
    # Try SKU exact match first
    if sku:
        sku_normalized = normalize_string(sku)
        for mat in materials:
            if mat.get('sku') and normalize_string(mat['sku']) == sku_normalized:
                return {
                    'material_id': mat['id'],
                    'material_name': mat['name'],
                    'material_sku': mat.get('sku'),
                    'confidence': 1.0,
                    'match_type': 'sku_exact'
                }
    
    # Try fuzzy match on description
    if description:
        desc_normalized = normalize_string(description)
        material_names = [normalize_string(mat['name']) for mat in materials]
        
        matches = get_close_matches(desc_normalized, material_names, n=1, cutoff=cutoff)
        if matches:
            # Find the material that matched
            matched_name = matches[0]
            for mat in materials:
                if normalize_string(mat['name']) == matched_name:
                    # Calculate actual similarity score
                    confidence = SequenceMatcher(None, desc_normalized, matched_name).ratio()
                    
                    return {
                        'material_id': mat['id'],
                        'material_name': mat['name'],
                        'material_sku': mat.get('sku'),
                        'confidence': confidence,
                        'match_type': 'description_fuzzy'
                    }
    
    return None


def parse_csv(file_content: bytes, column_mapping: Optional[Dict[str, str]] = None, delimiter: str = ',') -> Dict:
    """
    Parse CSV invoice file with flexible column mapping
    
    Args:
        file_content: CSV file content as bytes
        column_mapping: Optional mapping of column names, e.g.:
            {'sku': 'Cod', 'description': 'Denumire', 'quantity': 'Cantitate', ...}
        delimiter: CSV delimiter (default: ',', common Romanian: ';')
    
    Returns:
        Dict with extracted invoice data
    """
    try:
        # Decode content
        text_content = file_content.decode('utf-8-sig')  # Handle BOM
        
        # Try to detect delimiter if not specified
        sample = text_content[:1024]
        if delimiter == ',' and sample.count(';') > sample.count(','):
            delimiter = ';'
        
        # Parse CSV
        reader = csv.DictReader(io.StringIO(text_content), delimiter=delimiter)
        
        # Get fieldnames
        fieldnames = reader.fieldnames
        if not fieldnames:
            return {"error": "CSV has no headers", "items": []}
        
        # Default column mapping (common patterns - English and Romanian)
        default_mapping = {
            'sku': ['sku', 'cod', 'code', 'item_code', 'material_code'],
            'description': ['description', 'denumire', 'nume', 'name', 'produs', 'product'],
            'quantity': ['quantity', 'qty', 'cantitate', 'cant', 'cant.'],
            'unit': ['unit', 'um', 'u.m.', 'unitate'],
            'unit_price': ['unit_price', 'price', 'pret', 'pret_unitar', 'pretunit', 'pret_unit'],
            'total_price': ['total', 'total_price', 'sum', 'suma', 'valoare']
        }
        
        # Build actual mapping
        actual_mapping = {}
        if column_mapping:
            # Use provided mapping
            actual_mapping = column_mapping
        else:
            # Auto-detect columns
            for field_type, patterns in default_mapping.items():
                for fieldname in fieldnames:
                    fn_normalized = normalize_string(fieldname)
                    if any(normalize_string(pattern) in fn_normalized or fn_normalized in normalize_string(pattern) 
                           for pattern in patterns):
                        actual_mapping[field_type] = fieldname
                        break
        
        # Extract items
        items = []
        for row in reader:
            if not row:
                continue
            
            # Extract fields using mapping
            item = {}
            
            # SKU (optional)
            if 'sku' in actual_mapping and actual_mapping['sku'] in row:
                item['sku'] = row[actual_mapping['sku']].strip()
            
            # Description (required)
            if 'description' in actual_mapping and actual_mapping['description'] in row:
                item['description'] = row[actual_mapping['description']].strip()
            else:
                continue  # Skip rows without description
            
            # Quantity (required)
            if 'quantity' in actual_mapping and actual_mapping['quantity'] in row:
                try:
                    # Handle Romanian comma decimal separator
                    qty_str = row[actual_mapping['quantity']].strip().replace(',', '.')
                    item['quantity'] = float(qty_str)
                except (ValueError, AttributeError):
                    continue  # Skip invalid quantity
            else:
                continue  # Skip rows without quantity
            
            # Unit (optional)
            if 'unit' in actual_mapping and actual_mapping['unit'] in row:
                item['unit'] = row[actual_mapping['unit']].strip()
            
            # Unit price (optional)
            if 'unit_price' in actual_mapping and actual_mapping['unit_price'] in row:
                try:
                    # Handle Romanian comma decimal separator
                    price_str = row[actual_mapping['unit_price']].strip().replace(',', '.')
                    item['unit_price'] = float(price_str)
                except (ValueError, AttributeError):
                    pass
            
            # Total price (optional)
            if 'total_price' in actual_mapping and actual_mapping['total_price'] in row:
                try:
                    # Handle Romanian comma decimal separator
                    total_str = row[actual_mapping['total_price']].strip().replace(',', '.')
                    item['total_price'] = float(total_str)
                except (ValueError, AttributeError):
                    pass
            
            # Calculate missing values if possible
            if 'unit_price' in item and 'total_price' not in item:
                item['total_price'] = item['quantity'] * item['unit_price']
            elif 'total_price' in item and 'unit_price' not in item and item['quantity'] > 0:
                item['unit_price'] = item['total_price'] / item['quantity']
            
            items.append(item)
        
        # Calculate total
        total_amount = sum(item.get('total_price', 0) for item in items)
        
        return {
            "supplier": None,  # Not available in CSV
            "invoice_number": None,  # Not available in CSV
            "invoice_date": None,  # Not available in CSV
            "total_amount": total_amount if total_amount > 0 else None,
            "items": items,
            "detected_columns": list(actual_mapping.keys()),
            "column_mapping": actual_mapping
        }
    
    except Exception as e:
        logger.error(f"Error parsing CSV: {e}")
        return {"error": str(e), "items": []}
